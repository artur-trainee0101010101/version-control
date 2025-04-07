from docx import Document

# Data for the table
data = [
    ["Activity / Product", "Annual Output name", "Start date", "End date", 
     "Estimated cost w/o staff", "DID Int. FTEs", 
     "KPI affected (main benefit)", "CSFD ranking"],
    ["Opening up IPEP for candidate countries", 
     "Pilot with a selected third country (Ukraine)", "01/01/2025", "Early 2026", 
     "€185,000", "", "", "Increase in IPR enforcement", ""],
    ["Analysis (SRS) e-Notification on detentions", 
     "Digitalisation of detention workflow analysis", "01/01/2025", "31/12/2025", 
     "€50,000", "", "", "Efficiency gains", ""],
    ["Access levels and monitoring for enforcers", 
     "Create different access levels and monitoring tools", "01/01/2025", "31/12/2026", 
     "€238,000", "", "", "Enhanced enforcement control", ""],
    ["IPEP Data Hub", "Transition from annual report to interactive data hub", 
     "01/01/2025", "31/12/2025", "€60,000", "", "", "Broader data access", ""],
    ["Translations", "Support for new languages and intermediary adaptations", 
     "01/01/2025", "31/12/2025", "€20,000", "", "", "Broader enforcement scope", ""],
    ["Events", "IPEP Steering Group meetings (Customs, Police, Rights holders)", 
     "01/01/2025", "31/12/2025", "€47,412", "", "", "Enhanced collaboration", ""]
]

# Create a new Word Document
doc = Document()

# Add a title or heading (optional)
doc.add_heading('Project Activities and Outputs', level=1)

# Add a table to the document
table = doc.add_table(rows=1, cols=len(data[0]))

# Add headers
hdr_cells = table.rows[0].cells
for i, header in enumerate(data[0]):
    hdr_cells[i].text = header

# Add the data rows
for row_data in data[1:]:
    row_cells = table.add_row().cells
    for i, item in enumerate(row_data):
        row_cells[i].text = item

# Save the document
doc.save("Project_Activities_and_Outputs.docx")
